import os
from typing import Dict, List
import spacy
import pdfplumber
from docx import Document
from collections import Counter
import re

nlp = spacy.load("en_core_web_sm")

# Enhanced skill/tech extraction from resume PDF/DOCX
def resume_parser(resume_path: str) -> Dict:
    text = ""
    if resume_path.lower().endswith(".pdf"):
        with pdfplumber.open(resume_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    elif resume_path.lower().endswith(".docx"):
        doc = Document(resume_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    else:
        raise ValueError(f"Unsupported file format: {resume_path}")
    
    # Enhanced skill extraction with predefined tech skills
    
    # Enhanced skill extraction with predefined tech skills
    tech_skills = {
        'python', 'java', 'javascript', 'typescript', 'react', 'node.js', 'angular', 'vue',
        'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'terraform', 'ansible', 'jenkins',
        'git', 'linux', 'sql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
        'devops', 'ci/cd', 'microservices', 'api', 'rest', 'graphql', 'kafka',
        'machine learning', 'ai', 'data science', 'pandas', 'numpy', 'tensorflow',
        'pytorch', 'scikit-learn', 'spark', 'hadoop', 'tableau', 'powerbi'
    }
    
    doc = nlp(text)
    tokens = [t.text.lower() for t in doc if t.is_alpha and not t.is_stop]
    
    # Extract entities and tech skills
    entities = [ent.text.lower() for ent in doc.ents if ent.label_ in ("ORG", "PRODUCT", "PERSON", "GPE")]
    found_skills = []
    
    # Match tech skills in text
    text_lower = text.lower()
    for skill in tech_skills:
        if skill in text_lower:
            found_skills.append(skill)
    
    # Extract years of experience with regex
    experience_pattern = r'(\d+)[\+\s]*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)'
    experience_matches = re.findall(experience_pattern, text_lower)
    total_experience = max([int(x) for x in experience_matches], default=0)
    
    # Most common meaningful tokens
    common = [w for w, count in Counter(tokens).most_common(30) if len(w) > 2]
    
    all_skills = list(set(found_skills + entities + common[:15]))
    
    return {
        "text": text,
        "skills": all_skills,
        "tech_skills": found_skills,
        "tokens": tokens,
        "experience_years": total_experience,
        "entities": entities
    }
